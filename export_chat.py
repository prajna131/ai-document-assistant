"""
export_chat.py
--------------
Generates a complete "Project & Interview Guide" as a Word document (.docx).
Reading this single document should let you confidently explain the whole
project to an interviewer: what it is, how it works, what you built, the
challenges you solved, and answers to likely interview questions.

Run with:
    python -m pip install python-docx --trusted-host pypi.org --trusted-host files.pythonhosted.org
    python export_chat.py

Output: AI_Document_Assistant_Interview_Guide.docx  (in this folder)
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_code(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    return p


def add_bullets(doc, items):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def add_qa(doc, question, answer):
    p = doc.add_paragraph()
    p.add_run("Q: " + question).bold = True
    a = doc.add_paragraph()
    a.add_run("A: ").bold = True
    a.add_run(answer)


def main():
    doc = Document()

    # ============================================================ TITLE
    t = doc.add_heading("AI Document Assistant (RAG Chatbot)", level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = doc.add_paragraph("Project & Interview Guide")
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.runs[0].italic = True
    s.runs[0].bold = True
    doc.add_paragraph()

    # ============================================================ 1. PITCH
    doc.add_heading("1. Elevator Pitch (say this first)", level=1)
    doc.add_paragraph(
        "\"I built an AI Document Assistant - a web app where you upload your own "
        "documents (PDFs or text files) and chat with them in plain English. It "
        "uses Retrieval-Augmented Generation (RAG): instead of relying on the "
        "language model's memory, it searches your documents for the most "
        "relevant passages and feeds those to Google's Gemini model to generate "
        "an accurate, grounded answer - and it cites the exact source snippets. "
        "It's the same approach used by tools like ChatPDF and Perplexity. I "
        "built it in Python with LangChain, ChromaDB as the vector database, and "
        "Streamlit for the web interface.\""
    )

    # ============================================================ 2. WHAT/WHY
    doc.add_heading("2. What the Project Is & The Problem It Solves", level=1)
    doc.add_paragraph(
        "Large Language Models (LLMs) are powerful but have two problems: (1) they "
        "don't know about YOUR private documents, and (2) they can 'hallucinate' "
        "(make up facts). This project solves both by grounding the model in the "
        "user's own documents and showing the sources behind each answer."
    )
    doc.add_paragraph("Real use cases:")
    add_bullets(doc, [
        "Ask questions about a long research paper or textbook.",
        "Query company policy documents, manuals, or contracts.",
        "Summarise and cross-reference multiple documents at once.",
    ])

    # ============================================================ 3. WHAT IS RAG
    doc.add_heading("3. What is RAG? (explain it simply)", level=1)
    doc.add_paragraph(
        "RAG = Retrieval-Augmented Generation. It combines two steps:"
    )
    add_bullets(doc, [
        "Retrieval: find the most relevant pieces of text from a knowledge base "
        "using semantic (meaning-based) search.",
        "Generation: give those retrieved pieces + the user's question to an LLM, "
        "which writes the final answer using that context.",
    ])
    doc.add_paragraph(
        "Analogy: it's an open-book exam. Instead of answering from memory, the "
        "model first looks up the relevant pages, then answers using them."
    )

    # ============================================================ 4. FLOW
    doc.add_heading("4. How It Works - The Full Flow", level=1)

    doc.add_heading("Phase 1 - Ingestion (when a document is uploaded)", level=2)
    add_code(doc,
             "Upload file -> Extract text -> Split into chunks -> "
             "Create embeddings -> Store in ChromaDB")
    add_bullets(doc, [
        "Extract text: PyPDFLoader (PDF) or TextLoader (TXT) reads the raw text.",
        "Chunking: RecursiveCharacterTextSplitter cuts the text into ~1000-char "
        "overlapping pieces (overlap keeps sentences from being cut awkwardly).",
        "Embeddings: each chunk is sent to Gemini's embedding model, which returns "
        "a vector (a list of numbers capturing the chunk's meaning).",
        "Storage: the vectors + text are stored in ChromaDB, the vector database.",
    ])

    doc.add_heading("Phase 2 - Query (when a question is asked)", level=2)
    add_code(doc,
             "Question -> Embed question -> Search ChromaDB (MMR) -> "
             "Top chunks + Question -> LLM -> Answer + Sources")
    add_bullets(doc, [
        "The question is embedded into a vector too.",
        "ChromaDB finds the chunks whose vectors are closest in meaning "
        "(semantic search). We use MMR so results are relevant AND diverse.",
        "The retrieved chunks + the question are inserted into a prompt.",
        "Gemini generates the answer using only that context, and the app shows "
        "the source snippets it used.",
    ])

    # ============================================================ 5. CONCEPTS
    doc.add_heading("5. AI Concepts Demonstrated (with one-line explanations)", level=1)
    concepts = [
        ("LLM", "The model that generates answers (Google Gemini)."),
        ("RAG", "Retrieve relevant text, then generate a grounded answer."),
        ("Embeddings", "Turning text into vectors that capture meaning."),
        ("Vector Database", "ChromaDB stores vectors and searches them fast."),
        ("Semantic Search", "Finding text by meaning, not exact keyword match."),
        ("MMR", "Maximal Marginal Relevance - relevant + diverse retrieval."),
        ("Chunking", "Splitting documents into small searchable pieces."),
        ("Prompt Engineering", "Instructing the model to answer only from context."),
        ("Conversational Memory", "Passing chat history so follow-ups make sense."),
        ("Grounding / Citations", "Showing sources to reduce hallucinations."),
        ("LangChain", "Framework that wires all these pieces together."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"
    table.rows[0].cells[0].text = "Concept"
    table.rows[0].cells[1].text = "One-line explanation"
    for name, desc in concepts:
        row = table.add_row().cells
        row[0].text = name
        row[1].text = desc
    doc.add_paragraph()

    # ============================================================ 6. TECH STACK
    doc.add_heading("6. Tech Stack & Why Each Choice", level=1)
    add_bullets(doc, [
        "Python 3.11/3.12 - main language; 3.11/3.12 chosen because all AI "
        "libraries have prebuilt wheels (3.13 forced source builds that failed).",
        "LangChain - orchestrates the RAG pipeline (loaders, splitter, retriever, "
        "chains) so I don't reinvent the plumbing.",
        "Google Gemini - the LLM + embeddings; has a genuine free tier.",
        "ChromaDB - lightweight local vector database, no server needed.",
        "Streamlit - turns a Python script into a clean web UI very quickly.",
        "truststore - makes Python trust the OS certificate store (needed on "
        "networks that inspect SSL).",
    ])

    # ============================================================ 7. CODE
    doc.add_heading("7. Code Walkthrough (file by file)", level=1)
    add_bullets(doc, [
        "config.py - all settings in one place: API key loading, folder paths, "
        "model names, chunking sizes, retrieval settings; also enables truststore "
        "and disables ChromaDB telemetry.",
        "rag_engine.py - the core. Contains GeminiEmbeddings (custom embeddings "
        "class using Gemini's 'embedContent' with retry logic) and RAGEngine "
        "(ingest_file, query, and helpers). This is the file to know best.",
        "app.py - the Streamlit UI: file upload, 'Process Documents', the chat "
        "box, showing answers, source citations, and the loaded-documents list.",
        "list_models.py - a helper that prints which Gemini models your API key "
        "supports (used to fix 'model not found' errors).",
        "requirements.txt - the dependencies.",
    ])

    # ============================================================ 8. WHAT I DID
    doc.add_heading("8. What I Personally Built & Decided", level=1)
    add_bullets(doc, [
        "Designed the two-phase RAG pipeline (ingestion + query).",
        "Wrote a custom embeddings class when the library's built-in one called an "
        "API method the new embedding model didn't support.",
        "Chose MMR retrieval so answers can draw from multiple documents.",
        "Added source citations to make answers trustworthy and reduce hallucination.",
        "Optimised to a single LLM call per question to conserve the free-tier quota.",
        "Made the app robust: skips empty PDF pages, retries on network blips, and "
        "works on SSL-inspected networks.",
    ])

    # ============================================================ 9. CHALLENGES
    doc.add_heading("9. Real Challenges I Solved (great interview stories)", level=1)
    doc.add_paragraph(
        "Use the Problem -> Cause -> Fix format when telling these; it shows "
        "structured debugging."
    )
    challenges = [
        ("Corporate/college network blocked API calls (SSL error)",
         "The network inspects SSL with a custom certificate that Python's gRPC "
         "client didn't trust.",
         "Switched the client to REST transport and used the 'truststore' library "
         "so Python trusts the OS certificate store - same as the browser."),
        ("New embedding model rejected the library's API call",
         "The installed LangChain called 'batchEmbedContents', which the new "
         "'gemini-embedding-001' model doesn't support.",
         "Wrote a small custom embeddings class using the supported 'embedContent' "
         "method, one chunk at a time."),
        ("Model names kept returning 404 / 'not available'",
         "Model names change over time and differ per account.",
         "Wrote list_models.py to query available models, and used the "
         "'gemini-flash-latest' / 'gemini-flash-lite-latest' aliases that always "
         "point to current models."),
        ("Free-tier quota exceeded (HTTP 429)",
         "Each question originally made 2 LLM calls, and the newest model's free "
         "limit was tiny.",
         "Reduced to 1 LLM call per question, switched to a 'lite' model with a "
         "bigger free allowance, and used a fresh project for new quota."),
        ("Install failures on a different PC",
         "Python 3.13 had no prebuilt wheels for numpy / chroma-hnswlib, so pip "
         "tried to compile them and failed (no C++ compiler).",
         "Standardised on Python 3.11/3.12, which have prebuilt wheels - no "
         "compilation needed."),
        ("Scanned PDF added 0 chunks",
         "The PDF was an image (no selectable text), so there was nothing to "
         "embed.",
         "Diagnosed it as needing OCR; filtered out empty chunks so it fails "
         "gracefully instead of crashing."),
    ]
    for i, (p, c, f) in enumerate(challenges, 1):
        doc.add_heading(f"9.{i}  {p}", level=2)
        x = doc.add_paragraph(); x.add_run("Cause: ").bold = True; x.add_run(c)
        y = doc.add_paragraph(); y.add_run("Fix: ").bold = True; y.add_run(f)

    # ============================================================ 10. Q&A
    doc.add_heading("10. Likely Interview Questions & Strong Answers", level=1)
    qas = [
        ("What is RAG and why use it instead of just an LLM?",
         "RAG retrieves relevant text from a knowledge base and feeds it to the "
         "LLM. It lets the model answer about private/updated data it was never "
         "trained on, and it reduces hallucinations because answers are grounded "
         "in real retrieved text with citations."),
        ("What are embeddings?",
         "Numerical vectors that represent the meaning of text. Similar meanings "
         "produce vectors that are close together, which is what enables semantic "
         "search."),
        ("How does semantic search differ from keyword search?",
         "Keyword search matches exact words; semantic search matches meaning. "
         "'car' and 'automobile' are far apart for keywords but close in embedding "
         "space."),
        ("What is a vector database and why ChromaDB?",
         "It stores embeddings and finds the nearest ones quickly using similarity "
         "metrics. I chose ChromaDB because it's lightweight, runs locally, and "
         "needs no separate server."),
        ("Why chunk documents? How did you choose the size?",
         "LLMs and retrieval work better on small focused pieces. I used ~1000 "
         "characters with 150 overlap so context isn't lost at boundaries."),
        ("What is MMR and why did you use it?",
         "Maximal Marginal Relevance balances relevance with diversity, so the "
         "retrieved chunks aren't all near-duplicates - important when answering "
         "across multiple documents."),
        ("How do you reduce hallucinations?",
         "The system prompt tells the model to answer ONLY from the provided "
         "context and to say 'I don't know' otherwise, and the UI shows source "
         "snippets so answers are verifiable."),
        ("How does conversational memory work here?",
         "Previous messages are passed to the answer model as chat history, so "
         "follow-up questions like 'what about its price?' are understood."),
        ("How would you scale this to millions of documents?",
         "Use a managed/distributed vector DB (e.g., Pinecone, Milvus), batch "
         "embedding, metadata filtering, and caching; possibly add a re-ranking "
         "step after retrieval."),
        ("What are the limitations?",
         "It needs text (scanned PDFs need OCR), broad 'summarise everything' "
         "questions are weaker than specific ones, and free-tier rate limits "
         "apply. Improvements: OCR, streaming answers, re-ranking, and hybrid "
         "keyword+semantic search."),
    ]
    for q, a in qas:
        add_qa(doc, q, a)

    # ============================================================ 11. IMPROVE
    doc.add_heading("11. Limitations & Future Improvements (shows maturity)", level=1)
    add_bullets(doc, [
        "OCR support to read scanned/image PDFs.",
        "Streaming ('typing') answers for a better UX.",
        "A re-ranking model after retrieval for higher accuracy.",
        "Hybrid search (keyword + semantic).",
        "User accounts and per-user document collections.",
        "Deployment to the cloud (e.g., Streamlit Community Cloud).",
    ])

    # ============================================================ 12. RESUME
    doc.add_heading("12. Resume Bullet", level=1)
    doc.add_paragraph(
        "AI Document Assistant (RAG Chatbot) - Built a full-stack RAG web "
        "application enabling natural-language Q&A over user-uploaded documents. "
        "Implemented document chunking, vector embeddings, semantic search "
        "(ChromaDB), MMR retrieval, conversational memory, and grounded answer "
        "generation with source citations using LangChain and Google Gemini. "
        "Tech: Python, LangChain, ChromaDB, Streamlit, Gemini."
    )

    # ============================================================ 13. RUN
    doc.add_heading("13. How to Run (quick reference)", level=1)
    add_code(doc,
             "py -3.11 -m venv .venv\n"
             ".venv\\Scripts\\activate\n"
             "pip install -r requirements.txt\n"
             "# put GOOGLE_API_KEY=... in a file named .env\n"
             "python -m streamlit run app.py")

    out = "AI_Document_Assistant_Interview_Guide.docx"
    doc.save(out)
    print(f"\nSaved: {out}\n")


if __name__ == "__main__":
    main()
